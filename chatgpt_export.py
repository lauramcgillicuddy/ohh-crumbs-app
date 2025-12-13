"""
ChatGPT Export Parser
Parses ChatGPT JSON exports and allows viewing and exporting individual conversations
"""
import streamlit as st
import json
import zipfile
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT


def parse_conversations_json(json_data):
    """Parse the conversations.json file and extract all conversations"""
    try:
        conversations = json_data if isinstance(json_data, list) else [json_data]
        parsed_conversations = []

        for conv in conversations:
            conv_id = conv.get('id', 'unknown')
            title = conv.get('title', 'Untitled Conversation')
            create_time = conv.get('create_time')
            update_time = conv.get('update_time')
            mapping = conv.get('mapping', {})
            current_node = conv.get('current_node')

            # Extract messages from mapping
            messages = extract_messages(mapping, current_node)

            # Create a conversation object
            parsed_conv = {
                'id': conv_id,
                'title': title,
                'create_time': datetime.fromtimestamp(create_time) if create_time else None,
                'update_time': datetime.fromtimestamp(update_time) if update_time else None,
                'messages': messages,
                'message_count': len(messages)
            }
            parsed_conversations.append(parsed_conv)

        return parsed_conversations
    except Exception as e:
        st.error(f"Error parsing conversations: {str(e)}")
        return []


def extract_messages(mapping, current_node):
    """Extract messages from the mapping structure by traversing from current_node backwards"""
    messages = []

    if not mapping:
        return messages

    # Build the message chain by following parent references from current_node
    visited = set()
    node_id = current_node

    while node_id and node_id in mapping and node_id not in visited:
        visited.add(node_id)
        node = mapping[node_id]

        # Extract message data
        message_data = node.get('message')
        if message_data:
            author_role = message_data.get('author', {}).get('role', 'unknown')
            content = message_data.get('content', {})
            content_type = content.get('content_type', 'text')
            parts = content.get('parts', [])
            create_time = message_data.get('create_time')

            # Only add if there's actual content
            if parts and parts[0]:
                messages.append({
                    'role': author_role,
                    'content': '\n'.join([str(part) for part in parts if part]),
                    'timestamp': datetime.fromtimestamp(create_time) if create_time else None
                })

        # Move to parent
        node_id = node.get('parent')

    # Reverse to get chronological order
    messages.reverse()

    return messages


def export_to_word(conversation):
    """Export a conversation to a Word document"""
    doc = Document()

    # Add title
    title = doc.add_heading(conversation['title'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    if conversation['create_time']:
        p = doc.add_paragraph()
        p.add_run(f"Created: {conversation['create_time'].strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Spacer

    # Add messages
    for msg in conversation['messages']:
        role = msg['role']
        content = msg['content']

        # Add role header
        role_paragraph = doc.add_paragraph()
        role_run = role_paragraph.add_run(f"{role.upper()}")
        role_run.bold = True
        role_run.font.size = Pt(12)

        if role == 'user':
            role_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
        elif role == 'assistant':
            role_run.font.color.rgb = RGBColor(16, 163, 127)  # Green

        # Add timestamp if available
        if msg['timestamp']:
            time_run = role_paragraph.add_run(f" - {msg['timestamp'].strftime('%H:%M:%S')}")
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add content
        content_paragraph = doc.add_paragraph(content)
        content_paragraph.style = 'Normal'

        doc.add_paragraph()  # Spacer between messages

    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io


def export_to_pdf(conversation):
    """Export a conversation to a PDF document"""
    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='black',
        spaceAfter=30,
        alignment=TA_LEFT
    )

    user_style = ParagraphStyle(
        'UserStyle',
        parent=styles['Normal'],
        fontSize=11,
        textColor='#0066cc',
        fontName='Helvetica-Bold',
        spaceAfter=6
    )

    assistant_style = ParagraphStyle(
        'AssistantStyle',
        parent=styles['Normal'],
        fontSize=11,
        textColor='#10a37f',
        fontName='Helvetica-Bold',
        spaceAfter=6
    )

    content_style = ParagraphStyle(
        'ContentStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=12
    )

    # Add title
    elements.append(Paragraph(conversation['title'], title_style))

    # Add metadata
    if conversation['create_time']:
        meta_text = f"<i>Created: {conversation['create_time'].strftime('%Y-%m-%d %H:%M:%S')}</i>"
        elements.append(Paragraph(meta_text, styles['Normal']))

    elements.append(Spacer(1, 0.3 * inch))

    # Add messages
    for msg in conversation['messages']:
        role = msg['role']
        content = msg['content'].replace('<', '&lt;').replace('>', '&gt;')  # Escape HTML

        # Add role header
        if role == 'user':
            role_text = f"<b>USER</b>"
            elements.append(Paragraph(role_text, user_style))
        elif role == 'assistant':
            role_text = f"<b>ASSISTANT</b>"
            elements.append(Paragraph(role_text, assistant_style))
        else:
            role_text = f"<b>{role.upper()}</b>"
            elements.append(Paragraph(role_text, content_style))

        # Add content
        # Replace newlines with <br/> for PDF
        content_formatted = content.replace('\n', '<br/>')
        elements.append(Paragraph(content_formatted, content_style))
        elements.append(Spacer(1, 0.15 * inch))

    # Build PDF
    doc.build(elements)
    pdf_io.seek(0)

    return pdf_io


def show_chatgpt_export():
    """Main function to display the ChatGPT export parser interface"""
    st.title("💬 ChatGPT Export Parser")
    st.markdown("Upload your ChatGPT data export to view and export individual conversations")

    # Initialize session state
    if 'conversations' not in st.session_state:
        st.session_state.conversations = []
    if 'selected_conversation' not in st.session_state:
        st.session_state.selected_conversation = None

    # File upload section
    st.header("📁 Upload ChatGPT Export")
    uploaded_file = st.file_uploader(
        "Upload conversations.json or the exported ZIP file",
        type=['json', 'zip'],
        help="Export your ChatGPT data from Settings > Data controls > Export data"
    )

    if uploaded_file:
        try:
            # Handle ZIP files
            if uploaded_file.name.endswith('.zip'):
                with zipfile.ZipFile(uploaded_file) as zip_file:
                    # Look for conversations.json in the ZIP
                    conversations_file = None
                    for filename in zip_file.namelist():
                        if 'conversations.json' in filename:
                            conversations_file = filename
                            break

                    if conversations_file:
                        with zip_file.open(conversations_file) as json_file:
                            json_data = json.load(json_file)
                    else:
                        st.error("No conversations.json found in the ZIP file")
                        return
            else:
                # Handle direct JSON upload
                json_data = json.load(uploaded_file)

            # Parse conversations
            st.session_state.conversations = parse_conversations_json(json_data)
            st.success(f"✅ Loaded {len(st.session_state.conversations)} conversations!")

        except Exception as e:
            st.error(f"Error reading file: {str(e)}")
            return

    # Display conversations list
    if st.session_state.conversations:
        st.header("📚 Conversations")

        # Add filters and sorting
        col1, col2 = st.columns([2, 1])
        with col1:
            search_query = st.text_input("🔍 Search conversations", "")
        with col2:
            sort_by = st.selectbox("Sort by", ["Recent first", "Oldest first", "Title A-Z", "Most messages"])

        # Filter conversations
        filtered_convs = st.session_state.conversations
        if search_query:
            filtered_convs = [
                c for c in filtered_convs
                if search_query.lower() in c['title'].lower()
            ]

        # Sort conversations
        if sort_by == "Recent first":
            filtered_convs.sort(key=lambda x: x['update_time'] or x['create_time'] or datetime.min, reverse=True)
        elif sort_by == "Oldest first":
            filtered_convs.sort(key=lambda x: x['create_time'] or datetime.min)
        elif sort_by == "Title A-Z":
            filtered_convs.sort(key=lambda x: x['title'].lower())
        elif sort_by == "Most messages":
            filtered_convs.sort(key=lambda x: x['message_count'], reverse=True)

        # Display conversation list
        st.markdown(f"**Found {len(filtered_convs)} conversations**")

        for conv in filtered_convs:
            with st.expander(
                f"💬 {conv['title']} ({conv['message_count']} messages)",
                expanded=False
            ):
                col1, col2, col3 = st.columns([2, 1, 1])

                with col1:
                    if conv['create_time']:
                        st.caption(f"Created: {conv['create_time'].strftime('%Y-%m-%d %H:%M')}")
                    if conv['update_time']:
                        st.caption(f"Updated: {conv['update_time'].strftime('%Y-%m-%d %H:%M')}")

                with col2:
                    if st.button("👁️ View", key=f"view_{conv['id']}"):
                        st.session_state.selected_conversation = conv
                        st.rerun()

                with col3:
                    st.caption(f"ID: {conv['id'][:8]}...")

        # Display selected conversation
        if st.session_state.selected_conversation:
            st.divider()
            conv = st.session_state.selected_conversation

            st.header(f"💬 {conv['title']}")

            # Export buttons
            col1, col2, col3 = st.columns([1, 1, 3])
            with col1:
                word_doc = export_to_word(conv)
                st.download_button(
                    label="📄 Export to Word",
                    data=word_doc,
                    file_name=f"{conv['title'][:50]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="word_export"
                )

            with col2:
                pdf_doc = export_to_pdf(conv)
                st.download_button(
                    label="📕 Export to PDF",
                    data=pdf_doc,
                    file_name=f"{conv['title'][:50]}.pdf",
                    mime="application/pdf",
                    key="pdf_export"
                )

            st.divider()

            # Display messages
            for i, msg in enumerate(conv['messages']):
                if msg['role'] == 'user':
                    st.markdown(f"**🧑 USER**")
                    st.info(msg['content'])
                elif msg['role'] == 'assistant':
                    st.markdown(f"**🤖 ASSISTANT**")
                    st.success(msg['content'])
                else:
                    st.markdown(f"**{msg['role'].upper()}**")
                    st.text(msg['content'])

                if msg['timestamp']:
                    st.caption(f"🕐 {msg['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}")

                if i < len(conv['messages']) - 1:
                    st.markdown("---")

            # Back button
            if st.button("← Back to conversations list"):
                st.session_state.selected_conversation = None
                st.rerun()

    else:
        st.info("👆 Upload your ChatGPT export file to get started")

        # Help section
        with st.expander("ℹ️ How to export your ChatGPT data"):
            st.markdown("""
            1. Go to [ChatGPT](https://chat.openai.com/)
            2. Click on your profile in the bottom left
            3. Go to **Settings** > **Data controls**
            4. Click **Export data**
            5. You'll receive an email with a download link
            6. Download the ZIP file and upload it here

            The export contains a `conversations.json` file with all your chat history.
            """)
